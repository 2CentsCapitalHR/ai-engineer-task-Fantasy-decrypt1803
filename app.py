# app.py
import os
os.environ["USE_TF"] = "0"

import streamlit as st
import os
import tempfile
import json
import docx2txt
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import re
from typing import List, Dict, Tuple

# ---------------------------
# Configuration & resources
# ---------------------------
REFERENCE_FILES = [
    "reference_files/Data Sources.pdf",
    "reference_files/Task.pdf"
]

MODEL_NAME = "all-MiniLM-L6-v2"  # small fast embedding model

# Example checklists (extend as needed)
PROCESS_CHECKLISTS = {
    "company_incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Board Resolution",
        "Incorporation Application Form",
        "Register of Members and Directors",
    ],
    # add other processes here
}

# Keywords for doc type detection (simple, extendable)
DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa", "article of association"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Board Resolution": ["board resolution", "resolution of the board"],
    "Incorporation Application Form": ["application for incorporation", "incorporation application"],
    "Register of Members and Directors": ["register of members", "register of directors"],
    "UBO Declaration": ["ubo declaration", "ultimate beneficial owner"],
    # add more
}

# ---------------------------
# Utility functions
# ---------------------------

def extract_text_from_pdf(path: str) -> str:
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for p in range(len(reader.pages)):
            page = reader.pages[p]
            try:
                text += page.extract_text() + "\n"
            except Exception:
                pass
    return text

def extract_text_from_docx(path: str) -> str:
    return docx2txt.process(path) or ""

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    text = text.replace("\n", " ").strip()
    chunks = []
    i = 0
    while i < len(text):
        chunk = text[i:i+chunk_size]
        chunks.append(chunk)
        i += chunk_size - overlap
    return chunks

# ---------------------------
# RAG: build index from reference files
# ---------------------------
@st.cache_resource(show_spinner=False)
def build_reference_index(reference_paths: List[str], model_name: str = MODEL_NAME):
    model = SentenceTransformer(model_name)
    all_chunks = []
    metadata = []
    for path in reference_paths:
        if not os.path.exists(path):
            st.warning(f"Reference file not found: {path}")
            continue
        if path.lower().endswith(".pdf"):
            text = extract_text_from_pdf(path)
        elif path.lower().endswith(".docx") or path.lower().endswith(".doc"):
            text = extract_text_from_docx(path)
        else:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        chunks = chunk_text(text, chunk_size=800, overlap=150)
        for c in chunks:
            all_chunks.append(c)
            metadata.append({"source": os.path.basename(path)})
    if len(all_chunks) == 0:
        return {"model": model, "embeddings": None, "chunks": [], "meta": []}
    embeddings = model.encode(all_chunks, show_progress_bar=False, convert_to_numpy=True)
    return {"model": model, "embeddings": embeddings, "chunks": all_chunks, "meta": metadata}

def retrieve_relevant(reference_index, query: str, top_k: int = 3) -> List[Tuple[str, float]]:
    if reference_index["embeddings"] is None or len(reference_index["chunks"]) == 0:
        return []
    q_emb = reference_index["model"].encode([query], convert_to_numpy=True)
    sims = cosine_similarity(q_emb, reference_index["embeddings"])[0]
    idxs = np.argsort(sims)[::-1][:top_k]
    results = [(reference_index["chunks"][i], float(sims[i])) for i in idxs]
    return results

# ---------------------------
# Document classification / detection
# ---------------------------
def detect_doc_types_from_text(text: str) -> List[str]:
    text_l = text.lower()
    found = set()
    for doc_type, kwlist in DOC_TYPE_KEYWORDS.items():
        for kw in kwlist:
            if kw in text_l:
                found.add(doc_type)
                break
    return list(found)

def process_inference_from_uploaded_docs(uploaded_texts: Dict[str, str]) -> Tuple[str, List[str]]:
    """
    Very simple "which process?" logic: if AoA or MoA present -> company_incorporation
    """
    combined = " ".join(uploaded_texts.values()).lower()
    if "articles of association" in combined or "aoa" in combined or "memorandum of association" in combined:
        return "company_incorporation", list(uploaded_texts.keys())
    # Add further heuristics here
    return "unknown", list(uploaded_texts.keys())

# ---------------------------
# Red flag detection (simple rules)
# ---------------------------
def find_paragraph_indices_with_match(doc: Document, pattern: str) -> List[int]:
    indices = []
    for i, para in enumerate(doc.paragraphs):
        if re.search(pattern, para.text, flags=re.IGNORECASE):
            indices.append(i)
    return indices

def detect_red_flags(text: str, reference_index) -> List[Dict]:
    issues = []
    # 1) Jurisdiction check: look for "jurisdiction" or "court" lines
    jur_matches = re.findall(r'jurisdiction[^.]*\.|court[^.]*\.', text, flags=re.IGNORECASE)
    if jur_matches:
        selector = " ".join(jur_matches[:3])
        # retrieve ADGM guidance from reference index
        retrieved = retrieve_relevant(reference_index, selector, top_k=2)
        # If ADGM not explicitly referenced in the match -> flag
        if not re.search(r'adgm|ab[u]? dhabi global market', selector, flags=re.IGNORECASE):
            suggestion = "Update jurisdiction to explicitly reference ADGM Courts and ADGM Companies Regulations."
            citation = retrieved[0][0] if len(retrieved) else "See ADGM guidance in reference files."
            issues.append({
                "section_excerpt": selector,
                "issue": "Jurisdiction clause may not reference ADGM",
                "severity": "High",
                "suggestion": suggestion,
                "citation_snippet": citation
            })
    else:
        # if no jurisdiction lines found, that's a red flag
        suggestion = "Add an explicit jurisdiction clause specifying ADGM Courts if this document governs ADGM entities."
        retrieved = retrieve_relevant(reference_index, "jurisdiction ADGM", top_k=2)
        citation = retrieved[0][0] if len(retrieved) else "See ADGM guidance in reference files."
        issues.append({
            "section_excerpt": "",
            "issue": "Missing jurisdiction clause",
            "severity": "High",
            "suggestion": suggestion,
            "citation_snippet": citation
        })

    # 2) Missing signature block detection
    if not re.search(r'signature|signed by|for and on behalf|signature:', text, flags=re.IGNORECASE):
        issues.append({
            "section_excerpt": "",
            "issue": "Missing signature block or signatory section",
            "severity": "High",
            "suggestion": "Include explicit signature lines with name, role and date.",
            "citation_snippet": "ADGM templates typically have signatory sections; see reference files."
        })

    # 3) Ambiguous / non-binding language detection (example phrases)
    ambiguous_phrases = ["best endeavours", "reasonable endeavours", "to the best of my knowledge", "subject to", "may be requested"]
    for p in ambiguous_phrases:
        if re.search(re.escape(p), text, flags=re.IGNORECASE):
            issues.append({
                "section_excerpt": p,
                "issue": f"Ambiguous phrase: '{p}'",
                "severity": "Medium",
                "suggestion": f"Consider replacing '{p}' with specific obligations or measurable standards.",
                "citation_snippet": "Ambiguity reduces enforceability; prefer mandatory language where required."
            })

    # 4) Example: incorrect jurisdiction mention like 'UAE Federal Courts'
    if re.search(r'uae federal court|federal courts', text, flags=re.IGNORECASE):
        retrieved = retrieve_relevant(reference_index, "ADGM jurisdiction vs UAE Federal Courts", top_k=2)
        citation = retrieved[0][0] if len(retrieved) else "See ADGM guidance in reference files."
        issues.append({
            "section_excerpt": "mentions UAE Federal Courts",
            "issue": "Incorrect jurisdiction reference (UAE Federal Courts) for ADGM-governed documents",
            "severity": "High",
            "suggestion": "Replace references to UAE Federal Courts with ADGM Courts where the entity is ADGM-registered.",
            "citation_snippet": citation
        })

    return issues

# ---------------------------
# Annotate (insert inline review paragraphs into .docx)
# ---------------------------
def annotate_docx_with_issues(input_docx_path: str, output_docx_path: str, issues: List[Dict]):
    doc = Document(input_docx_path)
    # We will attempt to find paragraph matching each issue.section_excerpt (if provided)
    for issue in issues:
        excerpt = issue.get("section_excerpt", "").strip()
        inserted = False
        if excerpt:
            # try to locate a paragraph containing excerpt (case-insensitive)
            for i, para in enumerate(doc.paragraphs):
                if excerpt.lower() in para.text.lower():
                    # insert a new paragraph after this para
                    new_p = doc.add_paragraph()  # appended at end; we'll move down
                    # To place right after, we rebuild paragraphs: easier approach: append comment near end, but include the paragraph text position
                    # Simpler: append comment with reference to the excerpt text
                    new_run = new_p.add_run(f"REVIEW COMMENT (Auto): {issue['issue']} | Suggestion: {issue['suggestion']} | Citation: {issue['citation_snippet'][:200]}")
                    new_run.italic = True
                    new_run.font.size = Pt(10)
                    new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    inserted = True
                    break
        if not inserted:
            # Generic comment appended
            new_p = doc.add_paragraph()
            new_run = new_p.add_run(f"REVIEW COMMENT (Auto): {issue['issue']} | Suggestion: {issue['suggestion']} | Citation: {issue['citation_snippet'][:200]}")
            new_run.italic = True
            new_run.font.size = Pt(10)
            new_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    # Save annotated doc
    doc.save(output_docx_path)

# ---------------------------
# Main Streamlit app
# ---------------------------
def main():
    st.set_page_config(page_title="ADGM Corporate Agent Prototype", layout="wide")
    st.title("ADGM Corporate Agent — Prototype")
    st.markdown(
        "Upload `.docx` legal documents and the system will: detect document types, "
        "compare against ADGM checklists, flag common red flags, add inline review comments, "
        "and produce a JSON summary. (Prototype — not legal advice.)"
    )

    st.sidebar.header("References (RAG)")
    st.sidebar.write("Indexing reference files:")
    for p in REFERENCE_FILES:
        st.sidebar.write(f"- {p}")

    # Build the RAG index (cached)
    with st.spinner("Building reference index..."):
        reference_index = build_reference_index(REFERENCE_FILES)

    uploaded = st.file_uploader("Upload one or more .docx files", type=["docx"], accept_multiple_files=True)
    if uploaded:
        st.info("Processing uploaded documents...")
        uploaded_texts = {}
        tmp_input_paths = {}
        for f in uploaded:
            # Save to a temp file
            tf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tf.write(f.read())
            tf.flush()
            tf.close()
            tmp_input_paths[f.name] = tf.name
            text = extract_text_from_docx(tf.name)
            uploaded_texts[f.name] = text

        # detect types per doc
        doc_types = {}
        all_detected_types = set()
        for name, txt in uploaded_texts.items():
            detected = detect_doc_types_from_text(txt)
            doc_types[name] = detected
            for d in detected:
                all_detected_types.add(d)

        st.subheader("Detected document types")
        for name, dlist in doc_types.items():
            st.write(f"**{name}** → {dlist if dlist else 'Unknown / needs manual review'}")

        # Which process?
        process, uploaded_names = process_inference_from_uploaded_docs(uploaded_texts)
        st.markdown(f"**Inferred process:** `{process}`")

        # compare to checklist
        if process in PROCESS_CHECKLISTS:
            required = PROCESS_CHECKLISTS[process]
            present_types = list(all_detected_types)
            # simple matching by doc names / types
            missing = [r for r in required if not any(r.lower() in (t.lower() if isinstance(t,str) else "") for t in present_types)]
            st.subheader("Checklist verification")
            st.write(f"Required documents for `{process}`: {required}")
            st.write(f"Detected / Present: {present_types}")
            if missing:
                st.error(f"Missing required documents: {missing}")
            else:
                st.success("All required documents detected (based on simple keyword detection).")
        else:
            st.info("No checklist available for inferred process (or process unknown).")

        # Run red flag detection and annotate each file
        results = {"process": process, "documents_uploaded": len(uploaded), "required_documents": len(PROCESS_CHECKLISTS.get(process, [])), "missing_document": missing[0] if process in PROCESS_CHECKLISTS and missing else None, "issues_found": []}
        annotated_paths = {}

        for fname, input_path in tmp_input_paths.items():
            st.write(f"Analyzing `{fname}` ...")
            text = uploaded_texts[fname]
            issues = detect_red_flags(text, reference_index)
            # annotate docx
            outpath = os.path.join(tempfile.gettempdir(), f"annotated_{os.path.basename(fname)}")
            try:
                annotate_docx_with_issues(input_path, outpath, issues)
                annotated_paths[fname] = outpath
            except Exception as e:
                st.warning(f"Annotation failed for {fname}: {e}")
                annotated_paths[fname] = None

            # add to results
            for it in issues:
                results["issues_found"].append({
                    "document": fname,
                    "section": it.get("section_excerpt", ""),
                    "issue": it.get("issue", ""),
                    "severity": it.get("severity", ""),
                    "suggestion": it.get("suggestion", "")
                })

        # Show results and provide downloads
        st.subheader("Results summary")
        st.json(results)

        st.write("---")
        st.subheader("Download reviewed documents")
        for fname, path in annotated_paths.items():
            if path and os.path.exists(path):
                with open(path, "rb") as f:
                    btn = st.download_button(label=f"Download annotated {fname}", data=f.read(), file_name=f"annotated_{fname}")
            else:
                st.write(f"Annotated file not available for {fname}")

        # download JSON
        st.download_button("Download JSON report", data=json.dumps(results, indent=2), file_name="report.json", mime="application/json")

if __name__ == "__main__":
    main()
